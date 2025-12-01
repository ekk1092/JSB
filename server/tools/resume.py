import os
import json
import tempfile
import traceback
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import AzureOpenAI
from dotenv import load_dotenv

load_dotenv()

def get_azure_client():
    return AzureOpenAI(
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
    )

def add_section_heading(doc: Document, text: str):
    """Add a formatted section heading (compact for one-page layout)"""
    heading = doc.add_paragraph(text)
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(11)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(4)

def create_resume_docx(data: dict) -> str:
    """Create a professionally formatted resume in .docx format"""
    try:
        doc = Document()
        
        # Set tight margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Name
        name_para = doc.add_paragraph(data.get('name', ''))
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if name_para.runs:
            name_run = name_para.runs[0]
            name_run.font.size = Pt(16)
            name_run.font.bold = True
        name_para.paragraph_format.space_after = Pt(2)
        
        # Contact
        contact = data.get('contact', {})
        contact_parts = []
        if contact.get('email'): contact_parts.append(contact['email'])
        if contact.get('phone'): contact_parts.append(contact['phone'])
        if contact.get('location'): contact_parts.append(contact['location'])
        if contact.get('linkedin'): contact_parts.append(contact['linkedin'])
        
        if contact_parts:
            contact_para = doc.add_paragraph(' | '.join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if contact_para.runs:
                contact_para.runs[0].font.size = Pt(9)
            contact_para.paragraph_format.space_after = Pt(6)
        
        # Summary
        if data.get('summary'):
            add_section_heading(doc, "PROFESSIONAL SUMMARY")
            summary_para = doc.add_paragraph(data['summary'])
            summary_para.paragraph_format.space_after = Pt(6)
            for run in summary_para.runs:
                run.font.size = Pt(10)
        
        # Education
        if data.get('education'):
            add_section_heading(doc, "EDUCATION")
            for edu in data['education']:
                edu_para = doc.add_paragraph()
                edu_para.paragraph_format.space_after = Pt(1)
                degree_run = edu_para.add_run(edu.get('degree', ''))
                degree_run.bold = True
                degree_run.font.size = Pt(10)
                
                school_para = doc.add_paragraph(
                    f"{edu.get('school', '')} - {edu.get('location', '')} | {edu.get('graduation', '')}"
                )
                school_para.paragraph_format.space_after = Pt(4)
                if school_para.runs:
                    school_para.runs[0].font.size = Pt(9)
        
        # Experience
        if data.get('experience'):
            add_section_heading(doc, "EXPERIENCE")
            for i, exp in enumerate(data['experience']):
                title_para = doc.add_paragraph()
                title_para.paragraph_format.space_after = Pt(1)
                title_run = title_para.add_run(f"{exp.get('title', '')} - {exp.get('company', '')}")
                title_run.bold = True
                title_run.font.size = Pt(10)
                
                details_para = doc.add_paragraph()
                details_para.paragraph_format.space_after = Pt(2)
                details_run = details_para.add_run(
                    f"{exp.get('location', '')} | {exp.get('dates', '')}"
                )
                details_run.italic = True
                details_run.font.size = Pt(9)
                
                if exp.get('responsibilities'):
                    for resp in exp['responsibilities']:
                        bullet_para = doc.add_paragraph(resp, style='List Bullet')
                        bullet_para.paragraph_format.space_after = Pt(1)
                        bullet_para.paragraph_format.line_spacing = 1.0
                        for run in bullet_para.runs:
                            run.font.size = Pt(10)
                
                if i < len(data['experience']) - 1:
                    doc.add_paragraph().paragraph_format.space_after = Pt(4)
        
        # Skills
        if data.get('skills'):
            add_section_heading(doc, "SKILLS")
            skills_para = doc.add_paragraph(', '.join(data['skills']))
            for run in skills_para.runs:
                run.font.size = Pt(10)
        
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            return tmp.name
            
    except Exception as e:
        print(f"Error creating resume docx: {e}")
        traceback.print_exc()
        return None

def create_cover_letter_docx(data: dict) -> str:
    """Create a professionally formatted cover letter in .docx format"""
    try:
        doc = Document()
        
        # Set normal margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        # Applicant Name
        name_para = doc.add_paragraph(data.get('name', ''))
        if name_para.runs:
            name_para.runs[0].font.size = Pt(12)
            name_para.runs[0].font.bold = True
        name_para.paragraph_format.space_after = Pt(0)
        
        # Contact
        contact = data.get('contact', {})
        if contact.get('address'):
            doc.add_paragraph(contact['address']).paragraph_format.space_after = Pt(0)
        if contact.get('phone'):
            doc.add_paragraph(contact['phone']).paragraph_format.space_after = Pt(0)
        if contact.get('email'):
            doc.add_paragraph(contact['email']).paragraph_format.space_after = Pt(12)
            
        # Date
        if data.get('date'):
            doc.add_paragraph(data['date']).paragraph_format.space_after = Pt(12)
            
        # Recipient
        recipient = data.get('recipient', {})
        if recipient:
            for key in ['name', 'title', 'company', 'address']:
                if recipient.get(key):
                    doc.add_paragraph(recipient[key]).paragraph_format.space_after = Pt(0)
            doc.add_paragraph().paragraph_format.space_after = Pt(12) # Extra space after recipient block
            
        # Body
        if data.get('body_paragraphs'):
            for paragraph in data['body_paragraphs']:
                p = doc.add_paragraph(paragraph)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = 1.15
                
        # Closing
        doc.add_paragraph("Sincerely,").paragraph_format.space_after = Pt(36)
        doc.add_paragraph(data.get('name', ''))
        
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            return tmp.name
            
    except Exception as e:
        print(f"Error creating cover letter docx: {e}")
        traceback.print_exc()
        return None

def tailor_resume_tool(resume_text: str, job_description: str) -> str:
    """
    Tailors a resume and returns a JSON string with 'preview' (markdown) and 'file_path' (docx).
    """
    client = get_azure_client()
    deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4o")

    prompt = f"""
    You are an expert career coach and resume writer.
    
    JOB DESCRIPTION:
    {job_description}
    
    CURRENT RESUME:
    {resume_text}
    
    Task: Rewrite the resume to better match the job description.
    
    OUTPUT FORMAT:
    Return a JSON object with the following structure:
    {{
        "name": "Candidate Name",
        "contact": {{ "email": "...", "phone": "...", "location": "...", "linkedin": "..." }},
        "summary": "Professional summary...",
        "experience": [
            {{ "title": "...", "company": "...", "location": "...", "dates": "...", "responsibilities": ["...", "..."] }}
        ],
        "education": [
            {{ "degree": "...", "school": "...", "location": "...", "graduation": "..." }}
        ],
        "skills": ["...", "..."],
        "preview_markdown": "A brief markdown summary of the changes made and why."
    }}
    """

    response = client.chat.completions.create(
        model=deployment_name,
        messages=[
            {"role": "system", "content": "You are a helpful assistant that outputs JSON."},
            {"role": "user", "content": prompt}
        ],
        response_format={"type": "json_object"}
    )
    
    try:
        content = response.choices[0].message.content
        data = json.loads(content)
        
        # Generate DOCX
        file_path = create_resume_docx(data)
        
        result = {
            "preview": data.get("preview_markdown", "Resume tailored successfully."),
            "file_path": file_path
        }
        return json.dumps(result)
        
    except Exception as e:
        return json.dumps({"error": f"Failed to generate resume: {str(e)}"})

def generate_cover_letter_tool(resume_text: str, job_description: str) -> str:
    """
    Generates a cover letter and returns a JSON string with 'preview' (markdown) and 'file_path' (docx).
    """
    client = get_azure_client()
    deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4o")

    prompt = f"""
    You are an expert career coach.
    
    JOB DESCRIPTION:
    {job_description}
    
    RESUME:
    {resume_text}
    
    Task: Write a compelling cover letter.
    
    OUTPUT FORMAT:
    Return a JSON object with the following structure:
    {{
        "name": "Candidate Name",
        "contact": {{ "email": "...", "phone": "...", "address": "..." }},
        "date": "Month Day, Year",
        "recipient": {{ "name": "...", "title": "...", "company": "...", "address": "..." }},
        "body_paragraphs": ["Para 1...", "Para 2...", "Para 3..."],
        "preview_markdown": "A brief markdown summary of the cover letter strategy."
    }}
    """

    response = client.chat.completions.create(
        model=deployment_name,
        messages=[
            {"role": "system", "content": "You are a helpful assistant that outputs JSON."},
            {"role": "user", "content": prompt}
        ],
        response_format={"type": "json_object"}
    )
    
    try:
        content = response.choices[0].message.content
        data = json.loads(content)
        
        # Generate DOCX
        file_path = create_cover_letter_docx(data)
        
        result = {
            "preview": data.get("preview_markdown", "Cover letter generated successfully."),
            "file_path": file_path
        }
        return json.dumps(result)
        
    except Exception as e:
        return json.dumps({"error": f"Failed to generate cover letter: {str(e)}"})
