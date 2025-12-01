import streamlit as st
import asyncio
import os
import sys
from mcp import ClientSession, StdioServerParameters
from mcp.client.stdio import stdio_client
from openai import AzureOpenAI
from dotenv import load_dotenv
import pandas as pd
import io

# Load environment variables
load_dotenv()

# Page configuration
st.set_page_config(page_title="Job Assistant", page_icon="💼", layout="wide")

# Initialize Azure OpenAI Client
client = AzureOpenAI(
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
)
deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4o")

# Session State Initialization
if "messages" not in st.session_state:
    st.session_state.messages = []
if "resume_text" not in st.session_state:
    st.session_state.resume_text = None

# Sidebar for Resume Upload
with st.sidebar:
    st.title("📄 Resume Upload")
    uploaded_file = st.file_uploader("Upload your Resume (TXT/MD/PDF/DOCX)", type=["txt", "md", "pdf", "docx", "doc"])
    
    if uploaded_file:
        try:
            file_ext = uploaded_file.name.split(".")[-1].lower()
            
            if file_ext in ["txt", "md"]:
                stringio = io.StringIO(uploaded_file.getvalue().decode("utf-8"))
                st.session_state.resume_text = stringio.read()
            elif file_ext == "pdf":
                import pypdf
                pdf_reader = pypdf.PdfReader(uploaded_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                st.session_state.resume_text = text
            elif file_ext in ["docx", "doc"]:
                import docx
                doc = docx.Document(uploaded_file)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                st.session_state.resume_text = text
                
            st.success("Resume uploaded and processed successfully!")
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")

# Main Chat Interface
st.title("💼 AI Job Assistant")
st.markdown("I can help you search for jobs, tailor your resume, and write cover letters.")

# Display Chat History
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Async function to handle MCP interaction
async def run_chat_loop(user_input):
    # Determine connection method
    mcp_server_url = os.getenv("MCP_SERVER_URL")

    if mcp_server_url:
        from mcp.client.sse import sse_client
        client_context = sse_client(mcp_server_url)
    else:
        # Server parameters for local stdio
        server_params = StdioServerParameters(
            command=sys.executable,
            args=["server/main.py"],
            env=os.environ.copy()
        )
        client_context = stdio_client(server_params)

    async with client_context as (read, write):
        async with ClientSession(read, write) as session:
            # Initialize session
            await session.initialize()
            
            # List available tools
            tools = await session.list_tools()
            
            # Prepare tools for OpenAI
            openai_tools = []
            for tool in tools.tools:
                openai_tools.append({
                    "type": "function",
                    "function": {
                        "name": tool.name,
                        "description": tool.description,
                        "parameters": tool.inputSchema
                    }
                })

            # Append user message
            st.session_state.messages.append({"role": "user", "content": user_input})
            with st.chat_message("user"):
                st.markdown(user_input)

            # Prepare context for LLM
            messages = [
                {"role": "system", "content": f"You are a helpful job assistant. You have access to tools. If the user asks to tailor a resume or write a cover letter, use the uploaded resume text: {st.session_state.resume_text if st.session_state.resume_text else 'No resume uploaded yet.'}. When you generate a resume or cover letter using the tools, inform the user that they can download the file using the buttons that will appear below your response."}
            ] + st.session_state.messages

            # Call LLM
            with st.chat_message("assistant"):
                message_placeholder = st.empty()
                full_response = ""
                
                response = client.chat.completions.create(
                    model=deployment_name,
                    messages=messages,
                    tools=openai_tools,
                    tool_choice="auto"
                )
                
                response_message = response.choices[0].message
                
                # Handle Tool Calls
                if response_message.tool_calls:
                    messages.append(response_message)
                    
                    for tool_call in response_message.tool_calls:
                        function_name = tool_call.function.name
                        function_args = eval(tool_call.function.arguments) # Safe for internal tools, use json.loads in prod
                        
                        message_placeholder.markdown(f"*Calling tool: `{function_name}`...*")
                        
                        # Call MCP Tool
                        result = await session.call_tool(function_name, arguments=function_args)
                        content = str(result.content)

                        # Store generated content for download
                        if function_name in ["tailor_resume", "generate_cover_letter"]:
                            st.session_state.last_generated_content = content
                            st.session_state.last_generated_type = "resume" if function_name == "tailor_resume" else "cover_letter"
                        
                        # Add tool result to messages
                        messages.append({
                            "tool_call_id": tool_call.id,
                            "role": "tool",
                            "name": function_name,
                            "content": content
                        })
                    
                    # Get final response from LLM
                    second_response = client.chat.completions.create(
                        model=deployment_name,
                        messages=messages
                    )
                    full_response = second_response.choices[0].message.content
                else:
                    full_response = response_message.content

                message_placeholder.markdown(full_response)
                st.session_state.messages.append({"role": "assistant", "content": full_response})

                # Show download buttons if content was generated in this turn
                if "last_generated_content" in st.session_state and st.session_state.last_generated_content:
                    content = st.session_state.last_generated_content
                    doc_type = st.session_state.last_generated_type
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            label=f"Download {doc_type.replace('_', ' ').title()} (PDF)",
                            data=content,
                            file_name=f"{doc_type}.pdf",
                            mime="application/pdf"
                        )
                    with col2:
                        # Simple DOCX creation
                        import docx
                        doc = docx.Document()
                        doc.add_paragraph(content)
                        bio = io.BytesIO()
                        doc.save(bio)
                        st.download_button(
                            label=f"Download {doc_type.replace('_', ' ').title()} (DOCX)",
                            data=bio.getvalue(),
                            file_name=f"{doc_type}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    # Clear after showing (optional, or keep it?)
                    # st.session_state.last_generated_content = None

# User Input
if prompt := st.chat_input("What would you like to do?"):
    asyncio.run(run_chat_loop(prompt))